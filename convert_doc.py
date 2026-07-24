import os
import re
from docx import Document
from docx.shared import Inches

def create_doc():
    doc = Document()
    md_path = 'docs/USER_MANUAL.md'
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:].replace('**', ''), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].replace('**', ''), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].replace('**', ''), level=3)
        # Images
        elif line.startswith('> `![') or line.startswith('!['):
            match = re.search(r'\!\[.*?\]\((.*?)\)', line)
            if match:
                img_path = match.group(1)
                # resolve path
                img_path = img_path.replace('../Photos/', 'Photos/')
                if os.path.exists(img_path):
                    try:
                        doc.add_picture(img_path, width=Inches(5.5))
                    except Exception as e:
                        doc.add_paragraph(f"[Image error: {e}]")
                else:
                    doc.add_paragraph(f"[Image missing: {img_path}]")
        # Lists
        elif line.startswith('* ') or line.startswith('- '):
            doc.add_paragraph(line[2:].replace('**', ''), style='List Bullet')
        elif re.match(r'^\d+\.\s', line):
            text = line[line.find(' ')+1:].replace('**', '')
            doc.add_paragraph(text, style='List Number')
        # Quotes
        elif line.startswith('> '):
            doc.add_paragraph(line[2:].replace('**', ''), style='Quote')
        # HR
        elif line == '---':
            doc.add_paragraph('_' * 40)
        else:
            doc.add_paragraph(line.replace('**', ''))

    doc.save('docs/USER_MANUAL.docx')
    print("Docx created successfully.")

if __name__ == "__main__":
    create_doc()
